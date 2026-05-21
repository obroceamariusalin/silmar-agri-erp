from fastapi import FastAPI, Depends, HTTPException, File, UploadFile, status
from fastapi.security import HTTPBasic, HTTPBasicCredentials
from fastapi.staticfiles import StaticFiles
import openpyxl
from fastapi.responses import FileResponse 
import shutil
import os
from sqlalchemy.orm import Session
import models
import schemas
from database import engine, Base, SessionLocal
from fpdf import FPDF
import cloudinary
import cloudinary.uploader
import cloudinary.api
import secrets
import pdfplumber
import re
import io
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy import text
from datetime import datetime


Base.metadata.create_all(bind=engine)
os.makedirs("acte_salvate", exist_ok=True)

security = HTTPBasic()

def verifica_parola(credentials: HTTPBasicCredentials = Depends(security)):
    user_corect = secrets.compare_digest(credentials.username, "silmar")
    parola_corecta = secrets.compare_digest(credentials.password, "silmar")
    
    if not (user_corect and parola_corecta):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Date de conectare incorecte",
            headers={"WWW-Authenticate": "Basic"},
        )
    return credentials.username

app = FastAPI(dependencies=[Depends(verifica_parola)])

cloudinary.config( 
  cloud_name = "dmsgvmhy2", 
  api_key = "615547999745442", 
  api_secret = "2hSOqmsiV5gbUct_NJwSFxMDWQk",
  secure = True
)

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# --- RUTĂ PENTRU REPARAREA BAZEI DE DATE ---
@app.get("/repara-baza")
def repara_db(db: Session = Depends(get_db)):
    try:
        db.execute(text("ALTER TABLE arendasi ADD COLUMN nr_data_contract VARCHAR;"))
        db.commit()
    except Exception as e:
        pass # Ignorăm dacă coloana există deja
        
    try:
        db.execute(text("ALTER TABLE arendasi ADD COLUMN durata_contract VARCHAR DEFAULT '7';"))
        db.commit()
    except Exception as e:
        pass # Ignorăm dacă coloana există deja
        
    return {"Mesaj": "Baza de date a fost reparată! Acum poți salva numărul și durata contractului."}


# --- Ruta pentru a adăuga un Arendaș ---
@app.post("/arendasi/", response_model=schemas.ArendasResponse)
def adauga_arendas(arendas: schemas.ArendasCreate, db: Session = Depends(get_db)):

    if arendas.cnp and arendas.cnp.strip() != "":
        cnp_existent = db.query(models.Arendas).filter(models.Arendas.cnp == arendas.cnp).first()
        if cnp_existent:
            raise HTTPException(status_code=400, detail="Eroare: Acest CNP este deja inregistrat!")
    else:
        arendas.cnp = f"LIPSA-{secrets.token_hex(3).upper()}"
    
    noul_arendas = models.Arendas(
        nume_complet=arendas.nume_complet,
        cnp=arendas.cnp,
        adresa=arendas.adresa
    )
    
    db.add(noul_arendas)
    db.commit()
    db.refresh(noul_arendas) 
    
    return noul_arendas


 # --- Ruta pentru a afisa toti arendasii ---
@app.get("/arendasi/", response_model=list[schemas.ArendasResponse])
def citeste_arendasi(db: Session = Depends(get_db)):
    arendasi = db.query(models.Arendas).all()
    return arendasi

# --- Ruta pentru a adăuga un teren unui arendas ---
@app.post("/arendasi/{arendas_id}/terenuri/", response_model=schemas.TerenResponse)
def adauga_teren(arendas_id: int, teren: schemas.TerenCreate, db: Session = Depends(get_db)):
   
    om = db.query(models.Arendas).filter(models.Arendas.id == arendas_id).first()
    if not om:
        raise HTTPException(status_code=404, detail="Eroare: Arendasul nu a fost gasit!")

    noul_teren = models.Teren(
        comuna=teren.comuna,
        tarlaua=teren.tarlaua,       
        parcela=teren.parcela,       
        suprafata_ha=teren.suprafata_ha,
        arendas_id=arendas_id 
    )
    
   
    db.add(noul_teren)
    db.commit()
    db.refresh(noul_teren)
    
    return noul_teren


    # --- Ruta pentru incarcarea documentelor (Buletin, Act Teren, etc.) ---
@app.post("/arendasi/{arendas_id}/upload-document/")
def incarca_document(arendas_id: int, file: UploadFile = File(...)):

    nume_fisier_nou = f"{arendas_id}_{file.filename}"
    cale_salvare = f"acte_salvate/{nume_fisier_nou}"

    with open(cale_salvare, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    return {
        "mesaj": "Fisier incarcat si salvat cu succes!",
        "fisier": nume_fisier_nou
    }


# --- Înregistrarea unei plăți noi ---
@app.post("/arendasi/{arendas_id}/plati/", response_model=schemas.PlataResponse)
def inregistreaza_plata(arendas_id: int, plata: schemas.PlataCreate, db: Session = Depends(get_db)):
    noul_pachet = models.Plata(**plata.model_dump(), arendas_id=arendas_id)
    db.add(noul_pachet)
    db.commit()
    db.refresh(noul_pachet)
    return noul_pachet

# --- Calculul Bilanțului ---
@app.get("/arendasi/{arendas_id}/bilant/", response_model=schemas.BilantArendas)
def obtine_bilant(arendas_id: int, db: Session = Depends(get_db)):
    om = db.query(models.Arendas).filter(models.Arendas.id == arendas_id).first()
    if not om:
        raise HTTPException(status_code=404, detail="Arendasul nu exista")

    total_ha = sum(t.suprafata_ha for t in om.terenuri)

    total_datorat = total_ha * 600

    total_platit = sum(p.cantitate_kg for p in om.plati)
    
    return {
        "nume_arendas": om.nume_complet,
        "total_hectare": total_ha,
        "total_de_plata_kg": total_datorat,
        "total_achitat_kg": total_platit,
        "rest_de_plata_kg": total_datorat - total_platit,
        "istoric_plati": om.plati
    }

# ---  EDITARE ARENDAȘ ---
@app.patch("/arendasi/{arendas_id}", response_model=schemas.ArendasResponse)
def editeaza_arendas(arendas_id: int, date_noi: schemas.ArendasUpdate, db: Session = Depends(get_db)):
    om = db.query(models.Arendas).filter(models.Arendas.id == arendas_id).first()
    if not om:
        raise HTTPException(status_code=404, detail="Arendasul nu a fost gasit!")

    date_trimise = date_noi.model_dump(exclude_unset=True)

    for cheie, valoare in date_trimise.items():
        setattr(om, cheie, valoare)

    db.commit()
    db.refresh(om)
    return om

# --- EDITARE TEREN ---
@app.patch("/terenuri/{teren_id}", response_model=schemas.TerenResponse)
def editeaza_teren(teren_id: int, date_noi: schemas.TerenUpdate, db: Session = Depends(get_db)):
    teren = db.query(models.Teren).filter(models.Teren.id == teren_id).first()
    if not teren:
        raise HTTPException(status_code=404, detail="Terenul nu a fost gasit!")

    date_trimise = date_noi.model_dump(exclude_unset=True)

    for cheie, valoare in date_trimise.items():
        setattr(teren, cheie, valoare)

    db.commit()
    db.refresh(teren)
    return teren

# --- ȘTERGERE ARENDAȘ ---
@app.delete("/arendasi/{arendas_id}")
def sterge_arendas(arendas_id: int, db: Session = Depends(get_db)):
    om = db.query(models.Arendas).filter(models.Arendas.id == arendas_id).first()
    if not om:
        raise HTTPException(status_code=404, detail="Arendasul nu a fost gasit!")

    db.delete(om)
    db.commit()
    return {"mesaj": f"Arendasul {om.nume_complet} si toate terenurile lui au fost sterse cu succes!"}


def extrage_localitate(adresa: str):
    if not adresa:
        return ""

    match = re.search(r'(Comuna|Com\.|Mun\.|Municipiul|Ora[sș])\s+([A-ZĂÎÂȘȚa-zăîâșț\-\s]+)', adresa, re.IGNORECASE)
    if match:
        nume = match.group(2).split(',')[0].strip()
        prefix = match.group(1).title().replace('Com.', 'Comuna')
        return f"{prefix} {nume}"
    
    return adresa.split(',')[0].strip()

# --- RUTĂ PENTRU PLATĂ TOTALĂ (INDIVIDUALĂ SAU MULTIPLĂ) ---
@app.post("/arendasi/plata-totala-multipla/")
def plata_totala_multipla(date: dict, db: Session = Depends(get_db)):
    ids = date.get("ids", [])
    anul_agricol = date.get("an", 2026)
    
    if not ids:
        raise HTTPException(status_code=400, detail="Nu a fost selectat niciun arendaș.")

    for arendas_id in ids:
        om = db.query(models.Arendas).filter(models.Arendas.id == arendas_id).first()
        if not om:
            continue

        total_ha = sum(t.suprafata_ha for t in om.terenuri)
        total_datorat = total_ha * 600

        total_platit = sum(p.cantitate_kg for p in om.plati if p.anul_agricol == anul_agricol)
        
        rest_de_plata = total_datorat - total_platit

        if rest_de_plata > 0:
            pref = db.query(models.PreferintaAnuala).filter(
                models.PreferintaAnuala.arendas_id == arendas_id,
                models.PreferintaAnuala.anul_agricol == anul_agricol
            ).first()
            produs_ales = pref.tip_cereala if pref else "grau"

            noua_plata = models.Plata(
                arendas_id=arendas_id,
                anul_agricol=anul_agricol,
                cantitate_kg=rest_de_plata,
                produs=produs_ales,
                observatii="Plată totală automată (Bulk)"
            )
            db.add(noua_plata)

    db.commit()
    return {"mesaj": f"Plata a fost procesată cu succes pentru {len(ids)} arendași!"}

# --- EXPORT BORDEROU EXCEL ---
@app.get("/borderou/descarca-excel")
def descarca_borderou_excel(an: int = 2026, db: Session = Depends(get_db)):
    arendasi = db.query(models.Arendas).all()

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    ws_grau = wb.create_sheet(f"Grâu {an}")
    ws_porumb = wb.create_sheet(f"Porumb {an}")

    headers = [
        "NR. CRT.", "NUME SI PRENUME", "LOCALITATE", "SUPRAFATA HA", 
        "CANTIT. BRUTA KG", "CANT. DUPA DED. 40%", "IMPOZIT 10%", 
        "CANTITA NETA", "VALOARE TVA", "CANTITATE PRIMITA", "SEMNATURA"
    ]

    def setup_sheet(ws, titlu):
        ws.append(["S.C. SILMAR SOLUTION SRL"])
        ws.append(["C.F. 49659242"])
        ws.append([f"BORDEROU PRIVIND PLATA ARENDEI ({titlu}) PENTRU ANUL AGRICOL {an}"])
        ws.append([]) 
        ws.append(headers)

    setup_sheet(ws_grau, "GRÂU")
    setup_sheet(ws_porumb, "PORUMB")

    nr_crt_grau = 1
    nr_crt_porumb = 1

    for om in arendasi:
        suprafata_totala = sum(t.suprafata_ha for t in om.terenuri)

        if suprafata_totala == 0:
            continue

        pref = next((p for p in om.preferinte if p.anul_agricol == an), None)
        tip_cereala = pref.tip_cereala if pref else "grau"

        localitate_scurta = extrage_localitate(om.adresa)

        cant_bruta = round(suprafata_totala * 600)
        cant_dupa_ded = round(cant_bruta * 60 / 100)
        impozit = round(cant_dupa_ded * 10 / 100)
        cant_neta = cant_bruta - impozit
        valoare_tva = round(cant_neta * 11 / 100)
        cantitate_primita = cant_neta + valoare_tva

        rand_nou = [
            nr_crt_grau if tip_cereala == "grau" else nr_crt_porumb,
            om.nume_complet,
            localitate_scurta, 
            round(suprafata_totala, 3), 
            cant_bruta,
            cant_dupa_ded,
            impozit,
            cant_neta,
            valoare_tva,
            cantitate_primita,
            "" 
        ]

        if tip_cereala == "grau":
            ws_grau.append(rand_nou)
            nr_crt_grau += 1
        else:
            ws_porumb.append(rand_nou)
            nr_crt_porumb += 1

    nume_fisier = f"acte_salvate/Borderou_Arenda_{an}.xlsx"
    wb.save(nume_fisier)

    return FileResponse(
        path=nume_fisier, 
        filename=f"Borderou_Arenda_{an}.xlsx",
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


# --- EXPORT TABEL PRIMĂRIE ---
@app.get("/descarca-tabel-primarie")
def descarca_tabel_primarie(db: Session = Depends(get_db)):
    arendasi = db.query(models.Arendas).all()

    doc = Document()
    
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)

    p_antet = doc.add_paragraph()
    run_antet = p_antet.add_run("JUDETUL DOLJ\nCOMUNA ARGETOAIA\nPRIMARIA COMUNEI ARGETOAIA\n")
    run_antet.bold = True
    doc.add_paragraph(f"NR. .................... din {datetime.now().strftime('%d.%m.%Y')}")
   
    p_titlu = doc.add_paragraph()
    p_titlu.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titlu = p_titlu.add_run("TABEL CENTRALIZATOR AL CONTRACTELOR DE ARENDA\nSILMAR SOLUTION SRL, CUI 49659242 = arendaș")
    run_titlu.bold = True
    
    table = doc.add_table(rows=1, cols=13)
    table.style = 'Table Grid'
    
    capete_tabel = [
        "NR. CRT.", "NUME SI PRENUME ARENDATOR", "CNP", "NR./DATA CONTRACT DE ARENDA", 
        "DURATA DE ARENDARE (ANI)", "SUPRAFATA TOTALA (Ha)", "Arabil (ha)", "Pasuni", 
        "Fanete", "Vii", "Alte cat.", "Pruni", "TITULAR REGISTRUL AGRICOL"
    ]
    
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(capete_tabel):
        hdr_cells[i].text = text
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(8)
                
    total_general_ha = 0
    nr_crt = 1
    arendasi_sortati = sorted(arendasi, key=lambda x: x.nume_complet)
    
    for om in arendasi_sortati:
        suprafata_totala = sum(t.suprafata_ha for t in om.terenuri)
        if suprafata_totala == 0:
            continue
            
        total_general_ha += suprafata_totala
        
        row_cells = table.add_row().cells
        row_cells[0].text = str(nr_crt)
        row_cells[1].text = om.nume_complet
        row_cells[2].text = om.cnp if om.cnp and not om.cnp.startswith("LIPSA") else ".............."
        
        row_cells[3].text = om.nr_data_contract if om.nr_data_contract else "......./.........."
        row_cells[4].text = om.durata_contract if om.durata_contract else "7" 
        
        row_cells[5].text = f"{suprafata_totala:.4f}"
        row_cells[6].text = f"{suprafata_totala:.4f}"
        row_cells[7].text = "-"
        row_cells[8].text = "-"
        row_cells[9].text = "-"
        row_cells[10].text = "-"
        row_cells[11].text = "-"
        row_cells[12].text = om.nume_complet
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
        nr_crt += 1

    row_total = table.add_row().cells
    row_total[1].text = "TOTAL"
    row_total[5].text = f"{total_general_ha:.4f}"
    row_total[6].text = f"{total_general_ha:.4f}"
    
    for cell in row_total:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    
    doc.add_paragraph("\n")
    p_semn = doc.add_paragraph()
    p_semn.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_semn = p_semn.add_run("SECRETAR GENERAL,\nVÎNĂTORU DRAGOȘ-VICTOR")
    run_semn.bold = True

    nume_fisier = "acte_salvate/Tabel_Centralizator.docx"
    doc.save(nume_fisier)

    return FileResponse(
        path=nume_fisier, 
        filename=f"Tabel_Primarie_{datetime.now().year}.docx",
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


# --- CITIRE LISTĂ DOCUMENTE ---
@app.get("/arendasi/{arendas_id}/documente")
def get_documente(arendas_id: int):
    try:
        rezultate = cloudinary.api.resources(type="upload", prefix=f"{arendas_id}_", max_results=50)
        
        fisiere_gasite = []
        for doc in rezultate.get('resources', []):
            extensie = doc.get('format', '')
            nume_complet = f"{doc['public_id']}.{extensie}" if extensie else doc['public_id']
            
            fisiere_gasite.append({
                "nume": nume_complet,
                "url": doc['secure_url'] 
            })
        return fisiere_gasite
    except Exception as e:
        print(f"Eroare la citirea din cloud: {e}")
        return []


# ---  UPLOAD ÎN CLOUD ---
@app.post("/arendasi/{arendas_id}/upload/")
def upload_document(arendas_id: int, file: UploadFile = File(...)):
    nume_fara_extensie = file.filename.rsplit('.', 1)[0]
    nume_public = f"{arendas_id}_{nume_fara_extensie}"

    cloudinary.uploader.upload(file.file, public_id=nume_public, resource_type="auto")
    return {"mesaj": "Document salvat în siguranță în cloud!"}


# --- ȘTERGERE DIN CLOUD ---
@app.delete("/arendasi/{arendas_id}/documente/{nume_fisier}")
def sterge_document(arendas_id: int, nume_fisier: str):
    public_id = nume_fisier.rsplit('.', 1)[0]
    
    try:
        cloudinary.uploader.destroy(public_id, resource_type="image")
        cloudinary.uploader.destroy(public_id, resource_type="raw")
        return {"mesaj": "Document sters definitiv din cloud!"}
    except:
        raise HTTPException(status_code=404, detail="Eroare la stergere")

# --- ȘTERGERE TEREN INDIVIDUAL ---
@app.delete("/terenuri/{teren_id}")
def sterge_teren(teren_id: int, db: Session = Depends(get_db)):
    teren = db.query(models.Teren).filter(models.Teren.id == teren_id).first()
    if not teren:
        raise HTTPException(status_code=404, detail="Terenul nu a fost gasit!")
    
    db.delete(teren)
    db.commit()
    return {"mesaj": "Teren sters cu succes!"}

# --- ȘTERGERE PLATĂ ---
@app.delete("/plati/{id_plata}")
def sterge_plata(id_plata: int, db: Session = Depends(get_db)):
    plata = db.query(models.Plata).filter(models.Plata.id == id_plata).first()
    if not plata:
        raise HTTPException(status_code=404, detail="Plata nu a fost gasita")
    
    db.delete(plata)
    db.commit()
    return {"mesaj": "Plată ștearsă cu succes!"}

# --- EXTRAGERE DATE DIN PDF ---
@app.post("/extrage-contract/")
async def extrage_contract(file: UploadFile = File(...)):
    try:
        continut = await file.read()
        text_complet = ""
        
        with pdfplumber.open(io.BytesIO(continut)) as pdf:
            for page in pdf.pages:
                extragere = page.extract_text()
                if extragere:
                    text_complet += extragere + " "

        text = " ".join(text_complet.split())

        print(f"DEBUG TEXT PDF: {text}")

        nume_match = re.search(r'(?:Intre\s+doamna|Intre\s+domnul)\s+(.*?)\s+domiciliat', text, re.IGNORECASE)
        nume = nume_match.group(1).strip() if nume_match else ""

        adresa_match = re.search(r'domiciliat[aă]?\s+(?:in|în)\s+(.*?),\s+identificat', text, re.IGNORECASE)
        adresa = adresa_match.group(1).strip() if adresa_match else ""
        pattern_terenuri = r'Tarlaua\s*(\d+).*?parcel[aă]?\s*([\d\/A-Za-z]+).*?suprafa[tțţ]a\s*(?:de)?\s*([\d\s\.,]+)\s*(mp|ha)'
        
        terenuri_gasite = re.findall(pattern_terenuri, text, re.IGNORECASE)
        
        lista_terenuri = []
        for t in terenuri_gasite:
            tarla = t[0]
            parcela = t[1] 
            valoare_raw = t[2].replace(' ', '').replace(',', '.')
            unitate = t[3].lower()
            
            try:
                ha_parcela = float(valoare_raw)
                if unitate == 'mp':
                    ha_parcela = ha_parcela / 10000.0
            except:
                ha_parcela = 0.0
                
            lista_terenuri.append({
                "tarlaua": tarla,
                "parcela": parcela,
                "suprafata_ha": round(ha_parcela, 4)
            })

        return {
            "nume": nume,
            "adresa": adresa,
            "terenuri": lista_terenuri
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Eroare la citire PDF: {str(e)}")


# --- SALVARE PREFERINȚĂ CEREALE ---
@app.post("/arendasi/{arendas_id}/preferinte/")
def salveaza_preferinta(arendas_id: int, pref: schemas.PreferintaAnualaCreate, db: Session = Depends(get_db)):
    existenta = db.query(models.PreferintaAnuala).filter(
        models.PreferintaAnuala.arendas_id == arendas_id,
        models.PreferintaAnuala.anul_agricol == pref.anul_agricol
    ).first()
    
    if existenta:
        existenta.tip_cereala = pref.tip_cereala 
    else:
        noua = models.PreferintaAnuala(arendas_id=arendas_id, anul_agricol=pref.anul_agricol, tip_cereala=pref.tip_cereala)
        db.add(noua) 
        
    db.commit()
    return {"mesaj": "Preferința a fost actualizată!"}


@app.post("/jurnal/", response_model=schemas.JurnalCampResponse)
def adauga_jurnal(intrare: schemas.JurnalCampCreate, db: Session = Depends(get_db)):
    noua_intrare = models.JurnalCamp(**intrare.model_dump())
    db.add(noua_intrare)
    db.commit()
    db.refresh(noua_intrare)
    return noua_intrare

@app.get("/jurnal/", response_model=list[schemas.JurnalCampResponse])
def citeste_jurnal(db: Session = Depends(get_db)):
    return db.query(models.JurnalCamp).order_by(models.JurnalCamp.id.desc()).all()


@app.delete("/jurnal/{lucrare_id}")
def sterge_jurnal(lucrare_id: int, db: Session = Depends(get_db)):
    lucrare = db.query(models.JurnalCamp).filter(models.JurnalCamp.id == lucrare_id).first()
    if not lucrare:
        raise HTTPException(status_code=404, detail="Lucrarea nu a fost gasita")
    
    db.delete(lucrare)
    db.commit()
    return {"mesaj": "Lucrare ștearsă cu succes!"} 

@app.post("/genereaza-contract-rapid/")
def genereaza_contract_rapid(date: schemas.ContractRapidCreate):
    gen_text, domiciliat_text, identificat_text = "domnul/doamna", "domiciliat(ă)", "identificat(ă)"
    if date.cnp and len(date.cnp) > 0:
        if date.cnp[0] in ['1', '5', '7', '3']:
            gen_text, domiciliat_text, identificat_text = "domnul", "domiciliat", "identificat"
        elif date.cnp[0] in ['2', '6', '8', '4']:
            gen_text, domiciliat_text, identificat_text = "doamna", "domiciliată", "identificată"

    linii_terenuri = []
    total_ha = 0.0
    
    for t in date.terenuri:
        try:
            ha_val = float(t.ha.replace(',', '.'))
            total_ha += ha_val
            # Calculăm metrii pătrați
            suprafata_mp = int(ha_val * 10000) 
        except:
            suprafata_mp = 0
        linii_terenuri.append(f"- Tarlaua {t.tarla}, parcela {t.parcela}, în suprafață de {suprafata_mp} mp")
            
    text_terenuri = "\n".join(linii_terenuri) if linii_terenuri else "Nu a fost adăugat niciun teren."

    placeholders = {
        "{{GEN}}": gen_text,
        "{{NUME}}": date.nume.title(),
        "{{DOMICILIAT}}": domiciliat_text,
        "{{IDENTIFICAT}}": identificat_text,
        "{{adresa}}": date.adresa,
        "{{CNP}}": date.cnp,
        "{{SERIA}}": date.ci_seria.upper(),
        "{{NR_CI}}": date.ci_numar,
        "{{ELIBERAT}}": date.ci_eliberat.upper(),
        "{{DATA_CI}}": date.ci_data,
        "{{HA}}": f"{total_ha:.2f}",
        "{{LISTA_TERENURI}}": text_terenuri
    }

    doc = Document("template_contract.docx")

    for p in doc.paragraphs:
        for cheie, valoare in placeholders.items():
            if cheie in p.text:
                p.text = p.text.replace(cheie, valoare)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for cheie, valoare in placeholders.items():
                        if cheie in p.text:
                            p.text = p.text.replace(cheie, valoare)

    nume_fisier = "acte_salvate/Contract_Rapid_Generat.docx"
    doc.save(nume_fisier)

    return FileResponse(
        path=nume_fisier,
        filename=f"Contract_{date.nume.replace(' ', '_')}.docx",
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    
    
app.mount("/", StaticFiles(directory="../frontend", html=True), name="frontend")
